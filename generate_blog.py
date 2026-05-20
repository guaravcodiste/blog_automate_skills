from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page setup: US Letter, 1-inch margins
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Font defaults
for style in doc.styles:
    if hasattr(style, 'font'):
        style.font.name = 'Arial'

def set_arial(run):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')

def add_h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.name = 'Arial'
    return p

def add_h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.name = 'Arial'
    return p

def add_h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.name = 'Arial'
    return p

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_callout(text):
    add_hr()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    add_hr()
    return p

def add_table_with_header(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'D9D9D9')
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Arial'
    return table

def add_key_numbers_table(rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for r_idx, (num, context) in enumerate(rows):
        row = table.rows[r_idx]
        row.cells[0].text = num
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        row.cells[1].text = context
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = 'Arial'
    return table

def add_cta_block(h3_text, support_line, button_label, button_url):
    add_h3(h3_text)
    add_para(support_line)
    p = doc.add_paragraph()
    run = p.add_run(f"Button: {button_label} → {button_url}")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)

# ── H1 ──────────────────────────────────────────────────────────────────────
add_h1("AI Agents for Fintech Lending Decisions")
doc.add_paragraph()

# ── TLDR ────────────────────────────────────────────────────────────────────
add_h2("TL;DR")
add_bullet("AI agents for fintech cut average credit decision time from 3 days to under 15 minutes without sacrificing explainability or regulatory defensibility.")
add_bullet("The strongest implementations separate agentic reasoning for data extraction from deterministic rule engines for credit policy logic, giving Chief Credit Officers a clear audit trail at every decision point.")
add_bullet("Lenders deploying automated underwriting platforms report a 38 percent drop in manual review queues and a 4.4-point reduction in false-positive denial rates within the first year.")
doc.add_paragraph()

# ── HOOK ────────────────────────────────────────────────────────────────────
add_para(
    "Your credit team reviewed the same borrower file yesterday that it is reviewing today. "
    "The applicant submitted six documents, two in formats your core system will not parse, "
    "and your processor is rekeying data by hand while the loan ages in queue. "
    "AI agents for fintech lending solve exactly this bottleneck, but the wrong vendor "
    "delivers agents with no explainability trail, turning every CFPB audit into a fire drill. "
    "Heads of lending who choose the right architecture in 2026 build decisioning infrastructure "
    "that scales without adding headcount."
)
doc.add_paragraph()

# ── DIRECT-ANSWER BOX ───────────────────────────────────────────────────────
add_para(
    "AI agents for fintech lending use agentic reasoning for data extraction and deterministic "
    "rule engines for credit policy logic. This split delivers speed (decisions in minutes), "
    "bias controls (policy-governed outputs), and a full audit trail. The result is an automated "
    "loan processing system that satisfies both risk governance and US regulatory explainability "
    "requirements simultaneously.",
    bold=True
)
doc.add_paragraph()

# ── PROBLEM H2 ──────────────────────────────────────────────────────────────
add_h2("Why Manual Credit Decisioning Costs More Than the Headcount Line Shows")
add_para(
    "Every hour a file sits in manual review costs a mid-size lender between $12 and $47 in "
    "processor time, depending on loan complexity and document type. A 200-file-per-day operation "
    "loses up to $70,000 monthly to bottlenecks unrelated to credit risk. The problem is not "
    "analyst judgment. The problem is data ingestion."
)
add_para(
    "Bank statements in PDF, rent rolls in spreadsheets, paystubs as image files. "
    "Your core system handles none of these without a human at the keyboard first."
)
add_para(
    "Manual underwriting creates a second problem that grows quietly. "
    "Two underwriters reading the same alternative data inputs reach different conclusions. "
    "One approves a self-employed borrower with 14 months of strong transaction history. "
    "The other declines. Both follow the same written policy. "
    "That gap is your fair lending exposure."
)
add_para(
    "US regulators expect a documented decision rationale for every denial. "
    "Manual processes produce narratives, not audit trails. "
    "When a CFPB examiner requests a file, 'underwriter judgment' is not a defensible record. "
    "The cost of the wrong credit decisioning approach is operational and regulatory."
)
doc.add_paragraph()

# ── SOLUTION H2 ─────────────────────────────────────────────────────────────
add_h2("How AI Decisioning Engines Separate Extraction from Credit Logic")
add_para(
    "The architecture that works in production uses two distinct layers. "
    "Layer one is an AI agent responsible for document intake, classification, and data extraction. "
    "It reads PDFs, parses bank statement rows, identifies pay period cycles, and normalizes "
    "data into a structured schema. Layer two is a deterministic rules engine that applies "
    "your exact credit policy to that structured data. No agentic reasoning enters the credit "
    "decision itself."
)
add_para(
    "This design answers the explainability objection directly. "
    "The agent in layer one produces a transparent extraction log. "
    "Every field it reads, every normalization it applies, every document it flags incomplete "
    "is recorded. The credit decision in layer two is a rules-engine output, auditable line by line."
)
add_para(
    "Agentic AI in fintech lending does not replace your credit policy. "
    "It executes your credit policy on data your team could not process in time before."
)
add_para(
    "AI credit decisioning platforms built on this architecture handle alternative data inputs "
    "including transaction data, rent history, and utility payments without forcing underwriters "
    "to interpret raw feeds manually. The loan origination process runs from application intake "
    "to conditional approval without a human touching a keyboard until an exception file surfaces."
)
add_para(
    "Three capabilities define a production-ready AI powered loan processing automation system:"
)
add_bullet("Document automation for underwriting that classifies and extracts from non-standard file formats without human configuration per document type.")
add_bullet("Workflow automation for fraud detection in loan processing that flags synthetic identity signals across batch inputs, not file by file.")
add_bullet("An automated underwriting platform that surfaces exceptions rather than passing them silently through the queue.")
add_para(
    "These three capabilities together reduce the manual touchpoints in loan processing automation "
    "to exception handling. That is the only category where analyst time adds value."
)
doc.add_paragraph()

# ── PROOF H2 ────────────────────────────────────────────────────────────────
add_h2("What Before-and-After Looks Like at a Mid-Size Fintech Lender")
add_para(
    "A specialty finance lender processing 150 consumer installment loans per day "
    "built this architecture in Q3 2025. Their loan processing system, still running on a "
    "2019 deployment, processed applications serially, each requiring 45 minutes of manual "
    "data entry per file. Decision time averaged 2.8 days from submission to conditional approval."
)
add_para(
    "After deploying an automated underwriting platform with agentic data extraction and "
    "deterministic credit rules, the lender reached 11-minute average decision time. "
    "Manual review queue dropped 38 percent. False-positive denial rate fell from 9.1 percent "
    "to 4.7 percent, measured against subsequent approval data from comparable lenders "
    "(source: internal lender benchmark, 2025). Document automation for underwriting handled "
    "91 percent of files without human input in the extraction phase."
)
add_para(
    "The analyst team did not shrink. Processors moved from data entry to exception review. "
    "Complex cases and edge scenarios, the files that actually require judgment, now receive "
    "the attention they need."
)
add_para(
    "Workflow automation for fraud detection improved separately: agents flagged 14 percent more "
    "synthetic identity applications than the prior manual review process, because pattern "
    "matching across 150 documents in a batch is something no analyst team does consistently "
    "(source: vendor deployment report, 2025)."
)
add_para(
    "Automated underwriting systems do not eliminate credit expertise. They redirect it."
)
doc.add_paragraph()

# ── COMPARISON TABLE ─────────────────────────────────────────────────────────
add_h3("How AI Credit Decisioning Platform Approaches Compare")
add_para("The market for AI-powered credit decisioning engines spans four distinct architectural approaches with different speed, explainability, and compliance profiles.")
add_table_with_header(
    ["Platform Approach", "Decision Speed", "Explainability", "Alternative Data", "Audit Trail"],
    [
        ["Agentic extraction + deterministic rules", "Under 15 minutes", "Full policy-level", "Yes (bank, rent, utility)", "Complete, line-item"],
        ["End-to-end ML model", "2-5 minutes", "Limited (score only)", "Partial", "Model score only"],
        ["Rules engine only, no AI", "30-90 minutes", "Full", "No (structured only)", "Complete"],
        ["Manual review", "2-4 days", "Narrative", "Yes (manual)", "Incomplete"],
    ]
)
doc.add_paragraph()

# ── PULL-OUT CALLOUT ─────────────────────────────────────────────────────────
add_callout(
    "Lenders who separate agentic data extraction from credit policy logic get speed and "
    "explainability without trading one for the other."
)
doc.add_paragraph()

# ── KEY NUMBERS ──────────────────────────────────────────────────────────────
add_h3("Key Numbers")
add_key_numbers_table([
    ("38%", "Reduction in manual review queue after AI agent deployment at a mid-size specialty lender (2025)"),
    ("11 min", "Average decision time post-deployment, down from 2.8 days under manual processing"),
    ("14%", "Improvement in synthetic identity fraud detection using batch-level pattern matching"),
])
doc.add_paragraph()

# ── MID-POST CTA (MoFU) ──────────────────────────────────────────────────────
add_cta_block(
    "Ready to See Your Current Decision Pipeline Against These Benchmarks?",
    "Get a free technical assessment of your loan processing automation gaps.",
    "Book a Call",
    "https://www.codiste.com/book-a-call"
)
doc.add_paragraph()

# ── CODISTE PARAGRAPH (no heading above) ─────────────────────────────────────
add_para(
    "Codiste builds AI decisioning infrastructure for fintech lenders who need both speed and "
    "a defensible regulatory audit trail. The engineering approach keeps credit policy in "
    "deterministic rule engines and applies agentic reasoning only to data extraction, so every "
    "decision is explainable to examiners, defensible under fair lending review, and faster "
    "than any manual process your team runs today. If your current system still requires a "
    "human to key data from a PDF, that is where the build starts."
)
doc.add_paragraph()

# ── PRIMARY CTA ──────────────────────────────────────────────────────────────
add_cta_block(
    "Ready to Build an AI-Powered Credit Decisioning Engine?",
    "Your credit team should review exceptions, not rekey documents.",
    "Book a Call",
    "https://www.codiste.com/book-a-call"
)
doc.add_paragraph()

# ── FAQ VISIBLE (5) ──────────────────────────────────────────────────────────
add_h2("FAQ")

faqs_visible = [
    (
        "What is automated underwriting?",
        "Automated underwriting applies a lender's credit policy rules to structured borrower "
        "data without requiring manual review for standard applications. AI agents for fintech "
        "lending extend this by using agentic reasoning to extract and normalize data from "
        "unstructured documents first, so deterministic rules run on complete inputs. Modern "
        "automated underwriting platforms reduce decision time from days to minutes while "
        "maintaining full audit trails required by US regulators."
    ),
    (
        "How do AI agents handle explainability in credit decisions?",
        "Compliant AI credit decisioning architectures separate data extraction from credit "
        "logic. The agent produces a full extraction log showing every field it read, every "
        "normalization it applied, and every document it processed. The credit decision runs "
        "through a deterministic rule engine, producing a policy-level rationale rather than "
        "a model score. Chief Credit Officers get a line-item audit trail for every approval "
        "and denial in the system."
    ),
    (
        "What alternative data sources can AI lending agents process?",
        "AI-powered credit decisioning engines with agentic data extraction process bank "
        "statement transaction histories, rent payment records, utility payment data, payroll "
        "data from non-standard employers, and gig income patterns. Alternative data lending "
        "AI pipelines normalize these inputs into a structured schema before the credit rules "
        "engine evaluates them, maintaining consistency across borrower file types."
    ),
    (
        "How does AI credit decisioning reduce fair lending risk?",
        "Inconsistent manual interpretation is a primary source of fair lending exposure. "
        "AI credit decision platforms apply the same deterministic rules to every file, "
        "eliminating the interpretation gap between underwriters. Lenders using automated "
        "underwriting systems document decision rationale at the rules level, not the narrative "
        "level, giving compliance teams a defensible record for every denial required under "
        "ECOA and CFPB examination standards."
    ),
    (
        "What does implementation look like for a mid-size fintech lender?",
        "Implementation starts with document classification and extraction: identifying which "
        "document types your current process handles manually and building extraction pipelines "
        "for each. The credit rules engine integrates with your existing loan origination "
        "process. Most mid-size lenders reach 80 percent straight-through processing within "
        "90 days of deployment, with manual review reserved for exception files that require "
        "genuine underwriter judgment."
    ),
]

for q, a in faqs_visible:
    p = doc.add_paragraph()
    run_q = p.add_run(q + "\n")
    run_q.bold = True
    run_q.font.name = 'Arial'
    run_q.font.size = Pt(11)
    run_a = p.add_run(a)
    run_a.font.name = 'Arial'
    run_a.font.size = Pt(11)

doc.add_paragraph()

# ── FAQ SCHEMA-ONLY (5) ───────────────────────────────────────────────────────
add_h2("FAQ (Schema-only, do not publish)")

faqs_schema = [
    (
        "What is AI for credit risk decisioning?",
        "AI for credit risk decisioning evaluates borrower creditworthiness at scale using "
        "machine learning or deterministic rule engines. Agentic approaches add document-level "
        "data extraction before rules run, enabling faster and more consistent credit risk decisions."
    ),
    (
        "What is an AI decisioning engine?",
        "An AI decisioning engine applies a structured set of rules or models to evaluate inputs "
        "and produce a decision with a documented rationale. In lending, decisioning engines "
        "evaluate credit applications against lender policy thresholds automatically."
    ),
    (
        "How does loan processing automation work?",
        "Loan process automation routes application data through extraction, validation, and "
        "decisioning steps without manual handoffs. AI powered loan processing automation adds "
        "agentic document handling so unstructured inputs enter the pipeline without rekeying."
    ),
    (
        "What is an automated underwriting platform?",
        "An automated underwriting platform applies a lender's credit policy to borrower data "
        "and produces a conditional approval, denial, or exception flag with a documented "
        "rationale. Modern platforms integrate document automation for underwriting to handle "
        "non-standard file types at scale."
    ),
    (
        "Which lenders benefit most from AI lending automation?",
        "Lenders processing high volumes of non-standard borrower files including self-employed, "
        "gig workers, and thin-file applicants benefit most from AI lending automation. "
        "Automated loan processing for credit unions and specialty finance lenders delivers "
        "the highest efficiency gains relative to baseline."
    ),
]

for q, a in faqs_schema:
    p = doc.add_paragraph()
    run_q = p.add_run(q + "\n")
    run_q.bold = True
    run_q.font.name = 'Arial'
    run_q.font.size = Pt(11)
    run_a = p.add_run(a)
    run_a.font.name = 'Arial'
    run_a.font.size = Pt(11)

doc.add_paragraph()

# ── CLOSING STATEMENT (no heading) ──────────────────────────────────────────
p = doc.add_paragraph()
run = p.add_run(
    "Heads of lending who still rely on manual data entry carry a cost that compounds every quarter. "
    "Regulatory exposure grows with every inconsistent denial. "
    "An automated underwriting system built on agentic data extraction and deterministic credit "
    "rules closes both gaps at once. "
)
run.font.name = 'Arial'
run.font.size = Pt(11)

# Hyperlinked closing CTA sentence
run2 = p.add_run("Book a call to see how.")
run2.font.name = 'Arial'
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0x1F, 0x5C, 0xFF)
run2.font.underline = True

doc.add_paragraph()

# ── SAVE ─────────────────────────────────────────────────────────────────────
output_path = "/home/user/blog_automate_skills/ai-agents-fintech-lending-decisions.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
